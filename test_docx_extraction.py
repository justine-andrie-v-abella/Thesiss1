# ============================================================================
# test_docx_extraction.py  — run from project root:
#   python test_docx_extraction.py "path/to/your/questionnaire.docx"
# ============================================================================

import sys
import os
import re


def _reconstruct_matching_sections(text: str) -> str:
    lines  = text.splitlines()
    result = []
    i      = 0
    LETTERS = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'

    while i < len(lines):
        line = lines[i].strip()

        col_header = (
            re.search(r'column\s*a', line, re.I) and
            re.search(r'column\s*b', line, re.I)
        )

        if col_header:
            # Already inside a real [TABLE] block — keep as-is
            if result and result[-1].strip() == '[TABLE]':
                result.append(line)
                i += 1
                continue

            col_a_items = []
            col_b_items = []
            mixed_line  = None

            i += 1
            while i < len(lines):
                ln = lines[i].strip()
                if not ln:
                    i += 1
                    continue
                # Stop at next Roman-numeral section
                if re.match(r'^(II|III|IV|V|VI|VII|VIII|IX|X)[\.\s]', ln):
                    break
                if re.match(r'^bonus', ln, re.I):
                    break

                # Line with both a short term and a description separated by
                # spaces/tabs/pipes → last Column A item + first Column B item
                pipe_split = re.split(r'\s{3,}|\t|\s*\|\s*', ln, maxsplit=1)
                if len(pipe_split) == 2:
                    left, right = pipe_split[0].strip(), pipe_split[1].strip()
                    if left and right and len(left.split()) <= 5:
                        mixed_line = (left, right)
                        i += 1
                        continue

                words = ln.split()
                looks_like_term = (len(words) <= 6) and (ln[-1] not in '.?!')

                if looks_like_term and not col_b_items:
                    col_a_items.append(ln)
                else:
                    col_b_items.append(ln)

                i += 1

            if mixed_line:
                col_a_items.append(mixed_line[0])
                col_b_items.insert(0, mixed_line[1])

            if col_a_items and col_b_items:
                result.append("[TABLE]")
                result.append("Column A  |  Column B")
                max_rows = max(len(col_a_items), len(col_b_items))
                for idx in range(max_rows):
                    a_txt = col_a_items[idx] if idx < len(col_a_items) else ""
                    b_txt = col_b_items[idx] if idx < len(col_b_items) else ""
                    if a_txt and not a_txt[0].isdigit():
                        a_txt = f"{idx + 1}. {a_txt}"
                    if b_txt and (len(b_txt) < 2 or b_txt[1] not in ('.', ')')):
                        b_txt = f"{LETTERS[idx]}. {b_txt}"
                    result.append(f"{a_txt}  |  {b_txt}")
                result.append("[/TABLE]")
            else:
                result.extend(col_a_items)
                result.extend(col_b_items)

            continue

        result.append(line)
        i += 1

    return "\n".join(result)


def _extract_from_docx(file_path: str) -> str:
    import docx as _docx

    document = _docx.Document(file_path)
    parts    = []

    para_map  = {p._element: p for p in document.paragraphs}
    table_map = {t._element: t for t in document.tables}

    for child in document.element.body:
        tag = child.tag.split('}')[-1]

        if tag == 'p':
            para = para_map.get(child)
            if para and para.text.strip():
                parts.append(para.text.strip())

        elif tag == 'tbl':
            table = table_map.get(child)
            if table is None:
                continue

            grid = []
            for row in table.rows:
                seen  = set()
                cells = []
                for cell in row.cells:
                    txt     = cell.text.strip()
                    cell_id = id(cell._tc)
                    if cell_id not in seen:
                        seen.add(cell_id)
                        cells.append(txt)
                if any(cells):
                    grid.append(cells)

            if not grid:
                continue

            col_count   = max(len(r) for r in grid)
            is_matching = False

            if col_count == 2:
                data_rows = [r for r in grid if len(r) == 2 and r[0] and r[1]]
                for dr in data_rows:
                    if dr[0][:1].isdigit() or (dr[1][:1].isalpha() and dr[1][1:2] in ('.', ')')):
                        is_matching = True
                        break

            if is_matching:
                parts.append("[TABLE]")
                parts.append("Column A  |  Column B")
                for row in grid:
                    if len(row) == 2:
                        parts.append(f"{row[0]}  |  {row[1]}")
                    elif len(row) == 1 and row[0]:
                        parts.append(row[0])
                parts.append("[/TABLE]")
            else:
                for row in grid:
                    line = "  |  ".join(c for c in row if c)
                    if line:
                        parts.append(line)

    raw = "\n".join(parts)
    return _reconstruct_matching_sections(raw)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python test_docx_extraction.py <path_to_docx>")
        sys.exit(1)

    path = sys.argv[1]
    if not os.path.exists(path):
        print(f"File not found: {path}")
        sys.exit(1)

    result = _extract_from_docx(path)

    print("=" * 60)
    print("EXTRACTED TEXT OUTPUT")
    print("=" * 60)
    print(result)
    print("=" * 60)

    if "[TABLE]" in result:
        print("\n✅  Matching table detected!")
        print("\n--- TABLE BLOCK PREVIEW ---")
        in_table = False
        for ln in result.splitlines():
            if ln.strip() == "[TABLE]":
                in_table = True
            if in_table:
                print(ln)
            if ln.strip() == "[/TABLE]":
                in_table = False
    else:
        print("\n⚠️  No matching table detected.")