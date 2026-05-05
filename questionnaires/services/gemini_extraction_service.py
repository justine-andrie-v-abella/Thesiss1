# ============================================================================
# FILE: questionnaires/services/gemini_extraction_service.py
# ============================================================================

import json
import re
import time
import PyPDF2
import docx
import openpyxl
from typing import List, Dict
from django.conf import settings

# Transient HTTP codes worth retrying
_RETRYABLE = ('503', '429', 'UNAVAILABLE', 'RESOURCE_EXHAUSTED', 'rate limit', 'overloaded')


class GeminiQuestionnaireExtractor:
    """Extract or generate questions from uploaded files using Google Gemini API"""

    def __init__(self, api_key=None):
        self.api_key = api_key or getattr(settings, 'GEMINI_API_KEY', None)
        if not self.api_key:
            raise ValueError("GEMINI_API_KEY not found in settings")

        try:
            from google import genai
            self.client = genai.Client(api_key=self.api_key)
            self.model  = 'gemini-2.5-flash'
        except ImportError:
            raise ImportError("Please install google-genai: pip install google-genai")

    # =========================================================================
    # FILE EXTRACTION
    # =========================================================================

    def extract_text_from_file(self, file_path: str) -> str:
        """Route to the correct extractor based on file extension."""
        extension = file_path.lower().split('.')[-1]

        if extension == 'pdf':
            return self._extract_from_pdf(file_path)
        elif extension in ['docx', 'doc']:
            return self._extract_from_docx(file_path)
        elif extension in ['xlsx', 'xls']:
            return self._extract_from_excel(file_path)
        elif extension == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            raise ValueError(f"Unsupported file format: {extension}")

    def _extract_from_pdf(self, file_path: str) -> str:
        """
        Extract text from a PDF with formatting cues preserved.

        Uses PyMuPDF (fitz) as the primary engine because it:
          - Detects text colour → emits [RED:text] tags for reddish spans
          - Handles multi-column layouts correctly (sorts blocks by position)
          - Falls back to PyPDF2 if fitz is unavailable.
        """
        try:
            import fitz  # PyMuPDF
            return self._extract_from_pdf_pymupdf(file_path)
        except ImportError:
            pass
        # ── PyPDF2 fallback ────────────────────────────────────────────
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
            return text
        except Exception as e:
            raise Exception(f"Error extracting PDF: {str(e)}") from e

    def _extract_from_pdf_pymupdf(self, file_path: str) -> str:
        """
        PyMuPDF-based PDF extraction.

        Key features:
          1. Colour detection — spans with reddish colour get tagged [RED:text]
          2. Multi-column layout — text blocks are sorted by their vertical
             position within the same horizontal band so columns read top-to-
             bottom rather than being interleaved.
          3. Bold / italic detection via font flags.
        """
        import fitz

        COLUMN_BAND_TOLERANCE = 15  # px — spans within this y-delta share a row

        def is_reddish(color_int):
            """Return True for colours where R > 150, G < 100, B < 100."""
            if color_int is None:
                return False
            # fitz returns colour as a packed integer (ARGB on some builds; always sRGB)
            r = (color_int >> 16) & 0xFF
            g = (color_int >> 8)  & 0xFF
            b =  color_int        & 0xFF
            return r > 150 and g < 100 and b < 100

        def span_tag(span):
            """Return formatting tag for a span or None."""
            color = span.get('color', 0)
            if is_reddish(color):
                return 'RED'
            flags = span.get('flags', 0)
            # fitz font flags: bit 4 = italic, bit 3 = bold (serifed), bit 0 = superscript
            if flags & (1 << 4):   # italic
                return 'ITALIC'
            if flags & (1 << 3):   # bold (some fonts use bit 1 instead)
                return 'BOLD'
            if flags & (1 << 1):
                return 'BOLD'
            return None

        pages_text = []

        try:
            doc = fitz.open(file_path)
            for page in doc:
                page_width = page.rect.width

                # ── Collect all text spans with position info ──────────────
                raw_spans = []  # list of (x0, y0, x1, y1, text, tag)
                blocks = page.get_text('dict', flags=fitz.TEXT_PRESERVE_WHITESPACE)['blocks']
                for block in blocks:
                    if block.get('type') != 0:
                        continue
                    for line in block.get('lines', []):
                        for span in line.get('spans', []):
                            txt = span.get('text', '')
                            if not txt.strip():
                                continue
                            bbox = span.get('bbox', (0, 0, 0, 0))
                            tag  = span_tag(span)
                            raw_spans.append((bbox[0], bbox[1], bbox[2], bbox[3], txt, tag))

                if not raw_spans:
                    continue

                # ── Detect columns by looking for a clear x-gap ───────────
                # Sort by x0; if page has two clusters of x0, treat as 2-column
                x0_vals   = sorted(set(round(s[0] / 10) * 10 for s in raw_spans))
                mid_x     = page_width / 2

                # Group spans into left / right columns then sort each by y
                left_spans  = [s for s in raw_spans if s[0] < mid_x]
                right_spans = [s for s in raw_spans if s[0] >= mid_x]

                def spans_to_lines(spans):
                    """Cluster spans into logical lines by y0 proximity."""
                    if not spans:
                        return []
                    spans = sorted(spans, key=lambda s: (round(s[1] / 5) * 5, s[0]))
                    lines  = []
                    cur_y  = None
                    cur_line = []
                    for span in spans:
                        y = span[1]
                        if cur_y is None or abs(y - cur_y) > COLUMN_BAND_TOLERANCE:
                            if cur_line:
                                lines.append(cur_line)
                            cur_line = [span]
                            cur_y = y
                        else:
                            cur_line.append(span)
                    if cur_line:
                        lines.append(cur_line)
                    return lines

                def render_lines(lines):
                    result = []
                    for line in lines:
                        line_text = ''
                        for (x0, y0, x1, y1, txt, tag) in sorted(line, key=lambda s: s[0]):
                            stripped = txt.strip()
                            if not stripped:
                                line_text += txt
                                continue
                            if tag and not self._is_noise_run(stripped):
                                leading  = txt[: len(txt) - len(txt.lstrip())]
                                trailing = txt[len(txt.rstrip()):]
                                line_text += f'{leading}[{tag}:{stripped}]{trailing}'
                            else:
                                line_text += txt
                        rendered = line_text.strip()
                        if rendered:
                            result.append(rendered)
                    return result

                # If there are meaningful right-column spans, render columns separately
                if right_spans and len(right_spans) >= 2:
                    left_lines  = spans_to_lines(left_spans)
                    right_lines = spans_to_lines(right_spans)
                    page_lines  = render_lines(left_lines) + render_lines(right_lines)
                else:
                    all_lines  = spans_to_lines(raw_spans)
                    page_lines = render_lines(all_lines)

                pages_text.append('\n'.join(page_lines))

            doc.close()
            return '\n\n'.join(pages_text)

        except Exception as e:
            raise Exception(f"Error extracting PDF with PyMuPDF: {str(e)}") from e

    # =========================================================================
    # FORMATTING DETECTION HELPERS
    # =========================================================================

    @staticmethod
    def _get_run_format_tag(run):
        """
        Inspect a python-docx Run for answer-key formatting cues.
        Returns a tag string ('RED', 'HIGHLIGHT', 'UNDERLINE', 'BOLD', 'ITALIC')
        or None if the run has no special formatting.

        Priority order (most → least reliable as answer indicator):
          RED colour → HIGHLIGHT → UNDERLINE → BOLD → ITALIC
        """
        try:
            from docx.oxml.ns import qn
            rPr = run._r.find(qn('w:rPr'))
            if rPr is not None:

                # 1. Direct RGB colour -------------------------------------------
                color_el = rPr.find(qn('w:color'))
                if color_el is not None:
                    val = (color_el.get(qn('w:val')) or '').strip()
                    if val and val.lower() not in ('auto', '000000', 'ffffff') and len(val) == 6:
                        try:
                            r_v = int(val[0:2], 16)
                            g_v = int(val[2:4], 16)
                            b_v = int(val[4:6], 16)
                            # "Reddish" = high red channel, low green & blue
                            if r_v > 150 and g_v < 100 and b_v < 100:
                                return 'RED'
                        except ValueError:
                            pass

                # 2. Highlight ---------------------------------------------------
                hl = rPr.find(qn('w:highlight'))
                if hl is not None:
                    hl_val = (hl.get(qn('w:val')) or '').lower()
                    if hl_val and hl_val != 'none':
                        return 'HIGHLIGHT'

        except Exception:
            pass

        # 3. Underline (python-docx API is reliable here) -----------------------
        try:
            if run.underline:
                return 'UNDERLINE'
        except Exception:
            pass

        # 4. Bold (lower priority — headings are also bold) ---------------------
        try:
            if run.bold:
                return 'BOLD'
        except Exception:
            pass

        # 5. Italic -------------------------------------------------------------
        try:
            if run.italic:
                return 'ITALIC'
        except Exception:
            pass

        return None

    @staticmethod
    def _is_noise_run(text):
        """
        Returns True for short runs that are unlikely to be answer text even if
        formatted — e.g. question numbers ("1.", "2)"), pure punctuation, blanks.
        These are skipped to avoid tagging noise as answers.
        """
        stripped = text.strip()
        if not stripped:
            return True
        # Purely numeric/punctuation runs (question numbers, dashes, underscores)
        if re.match(r'^[\d\s\.\)\-\_]+$', stripped):
            return True
        # Very short runs that are just punctuation
        if len(stripped) <= 1 and not stripped.isalpha():
            return True
        return False

    def _extract_para_with_formatting(self, para):
        """
        Extract paragraph text, wrapping formatted runs with tags so Gemini
        can identify them as answer-key markers.
        e.g.  "1. [RED:Malware] is a type of malicious software."
        """
        parts = []
        for run in para.runs:
            text = run.text
            if not text:
                continue
            tag = self._get_run_format_tag(run)
            # Only tag runs that are meaningful (not question numbers / punctuation)
            if tag and not self._is_noise_run(text):
                stripped = text.strip()
                leading  = text[: len(text) - len(text.lstrip())]
                trailing = text[len(text.rstrip()):]
                parts.append(f"{leading}[{tag}:{stripped}]{trailing}")
            else:
                parts.append(text)
        result = ''.join(parts).strip()
        # Fallback: if all runs are plain, use para.text (handles edge cases)
        return result if result else para.text.strip()

    # =========================================================================
    # FILE EXTRACTION
    # =========================================================================

    def _extract_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX in reading order.
        Handles both real Word tables and paragraph-based column layouts.
        Preserves answer-key formatting cues ([RED:...], [BOLD:...], etc.)
        After raw extraction, runs _reconstruct_matching_sections() to
        reformat split Column A / Column B blocks into [TABLE] format.
        """
        try:
            document = docx.Document(file_path)
            parts    = []

            para_map  = {p._element: p for p in document.paragraphs}
            table_map = {t._element: t for t in document.tables}

            for child in document.element.body:
                tag = child.tag.split('}')[-1]

                # ── Plain paragraph ───────────────────────────────────────
                if tag == 'p':
                    para = para_map.get(child)
                    if para and para.text.strip():
                        parts.append(self._extract_para_with_formatting(para))

                # ── Real Word table ───────────────────────────────────────
                elif tag == 'tbl':
                    table = table_map.get(child)
                    if table is None:
                        continue

                    grid = []
                    for row in table.rows:
                        seen  = set()
                        cells = []
                        for cell in row.cells:
                            cell_id = id(cell._tc)
                            if cell_id not in seen:
                                seen.add(cell_id)
                                # Extract cell text with formatting cues
                                cell_parts = []
                                for para in cell.paragraphs:
                                    pt = self._extract_para_with_formatting(para)
                                    if pt:
                                        cell_parts.append(pt)
                                txt = ' '.join(cell_parts).strip()
                                cells.append(txt)
                        if any(cells):
                            grid.append(cells)

                    if not grid:
                        continue

                    col_count   = max(len(r) for r in grid)
                    is_matching = False

                    if col_count == 2:
                        data_rows = [r for r in grid if len(r) == 2 and r[0] and r[1]]
                        for dr in data_rows:
                            left_num  = dr[0][:1].isdigit()
                            right_let = dr[1][:1].isalpha() and dr[1][1:2] in ('.', ')')
                            if left_num or right_let:
                                is_matching = True
                                break

                    if is_matching:
                        parts.append("[TABLE]")
                        parts.append("Column A  |  Column B")
                        for row in grid:
                            if len(row) == 2:
                                parts.append(f"{row[0]}  |  {row[1]}")
                            elif len(row) == 1 and row[0]:
                                parts.append(row[0])
                        parts.append("[/TABLE]")
                    else:
                        for row in grid:
                            line = "  |  ".join(c for c in row if c)
                            if line:
                                parts.append(line)

            raw_text = "\n".join(parts)
            return self._reconstruct_matching_sections(raw_text)

        except Exception as e:
            raise Exception(f"Error extracting DOCX: {str(e)}") from e

    def _reconstruct_matching_sections(self, text: str) -> str:
        """
        Detects matching sections formatted as two separate paragraph blocks
        (all Column A items listed, then all Column B items listed) and
        stitches them into a proper [TABLE] block that Gemini can read.
        """
        lines   = text.splitlines()
        result  = []
        i       = 0
        LETTERS = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'

        while i < len(lines):
            line = lines[i].strip()

            col_header = (
                re.search(r'column\s*a', line, re.I) and
                re.search(r'column\s*b', line, re.I)
            )

            if col_header:
                # Already inside a real [TABLE] block — keep as-is
                if result and result[-1].strip() == '[TABLE]':
                    result.append(line)
                    i += 1
                    continue

                col_a_items = []
                col_b_items = []
                mixed_line  = None

                i += 1
                while i < len(lines):
                    ln = lines[i].strip()
                    if not ln:
                        i += 1
                        continue

                    # Stop at the next Roman-numeral section
                    if re.match(r'^(II|III|IV|V|VI|VII|VIII|IX|X)[\.\s]', ln):
                        break
                    if re.match(r'^bonus', ln, re.I):
                        break

                    # Line that has both a short term AND a long description
                    pipe_split = re.split(r'\s{3,}|\t|\s*\|\s*', ln, maxsplit=1)
                    if len(pipe_split) == 2:
                        left, right = pipe_split[0].strip(), pipe_split[1].strip()
                        if left and right and len(left.split()) <= 5:
                            mixed_line = (left, right)
                            i += 1
                            continue

                    words           = ln.split()
                    looks_like_term = (len(words) <= 6) and (ln[-1] not in '.?!')

                    if looks_like_term and not col_b_items:
                        col_a_items.append(ln)
                    else:
                        col_b_items.append(ln)

                    i += 1

                if mixed_line:
                    col_a_items.append(mixed_line[0])
                    col_b_items.insert(0, mixed_line[1])

                if col_a_items and col_b_items:
                    result.append("[TABLE]")
                    result.append("Column A  |  Column B")
                    max_rows = max(len(col_a_items), len(col_b_items))
                    for idx in range(max_rows):
                        a_txt = col_a_items[idx] if idx < len(col_a_items) else ""
                        b_txt = col_b_items[idx] if idx < len(col_b_items) else ""
                        if a_txt and not a_txt[0].isdigit():
                            a_txt = f"{idx + 1}. {a_txt}"
                        if b_txt and (len(b_txt) < 2 or b_txt[1] not in ('.', ')')):
                            b_txt = f"{LETTERS[idx]}. {b_txt}"
                        result.append(f"{a_txt}  |  {b_txt}")
                    result.append("[/TABLE]")
                else:
                    result.extend(col_a_items)
                    result.extend(col_b_items)

                continue

            result.append(line)
            i += 1

        return "\n".join(result)

    def _extract_from_excel(self, file_path: str) -> str:
        try:
            workbook = openpyxl.load_workbook(file_path)
            text = ""
            for sheet in workbook.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    text += " ".join([str(cell) for cell in row if cell]) + "\n"
            return text
        except Exception as e:
            raise Exception(f"Error extracting Excel: {str(e)}") from e

    # =========================================================================
    # AI PROCESSING
    # =========================================================================

    def extract_questions_with_ai(
        self,
        content: str,
        question_types: List[str],
        mode: str = 'extract',
        num_questions: int = 10,
    ) -> Dict:
        if mode == 'generate':
            prompt = self._build_generation_prompt(content, question_types, num_questions)
        else:
            prompt = self._build_extraction_prompt(content, question_types)

        max_retries = 3
        base_delay  = 3   # seconds; doubles each attempt

        last_error = None
        for attempt in range(max_retries):
            try:
                response = self.client.models.generate_content(
                    model=self.model,
                    contents=prompt,
                )

                response_text = response.text

                if '```json' in response_text:
                    response_text = response_text.split('```json')[1].split('```')[0]
                elif '```' in response_text:
                    response_text = response_text.split('```')[1].split('```')[0]

                response_text = response_text.strip()
                data = json.loads(response_text)

                if 'questions' not in data or not data['questions']:
                    raise Exception("AI response did not contain any questions")

                return data

            except json.JSONDecodeError as e:
                raise Exception(
                    f"Failed to parse AI response as JSON: {str(e)}\n"
                    f"Response preview: {response_text[:200]}"
                )
            except Exception as e:
                err_str    = str(e)
                is_transient = any(code in err_str for code in _RETRYABLE)

                if is_transient and attempt < max_retries - 1:
                    wait = base_delay * (2 ** attempt)   # 3s → 6s → 12s
                    print(f"[Gemini] transient error on attempt {attempt + 1}, "
                          f"retrying in {wait}s… ({err_str[:120]})")
                    time.sleep(wait)
                    last_error = e
                    continue

                # Non-retryable or exhausted retries — raise clean message
                raise Exception(err_str) from e

        raise Exception(str(last_error))

    # =========================================================================
    # PROMPTS
    # =========================================================================

    def _build_extraction_prompt(self, content: str, question_types: List[str]) -> str:
        return f"""You are a question scanner. Your ONLY job is to find and copy questions ALREADY WRITTEN in the text below.

STRICT RULES:
- DO NOT create, generate, invent, or add any new questions.
- ONLY include questions that are literally present in the provided text.
- If no questions are found, return {{"questions": []}}.
- Copy the question stem EXACTLY — but when an answer is embedded inline
  (via formatting tags [RED:…] / [BOLD:…] or a dash pattern), extract that
  answer into correct_answer and write the question WITHOUT the inline answer
  (replace it with a blank _______ if needed).
- NEVER leave correct_answer empty when you can see the answer in the text.

═══════════════════════════════════════════════════════
SECTION RECOGNITION
═══════════════════════════════════════════════════════
Recognize ALL of these as question sections:
- Roman numerals: I., II., III., IV., V., etc.
- Labels like: Multiple Choice, True or False, Identification, Matching Type,
  Fill in the Blank, Essay, Scenario-Based, Situation-Based, Spot the Error,
  Enumeration, Bonus, Short Answer, Problem Solving, etc.
- Numbered sub-questions inside any section are individual questions.

═══════════════════════════════════════════════════════
FORMATTED TEXT — ANSWER KEY DETECTION
═══════════════════════════════════════════════════════
The extracted text contains formatting tags where the teacher used special
formatting to mark the answer key. Use these to set correct_answer:

  [RED:text]       → red-coloured text  ← STRONGEST signal; almost always the answer
  [HIGHLIGHT:text] → highlighted text   ← very strong signal; likely the answer
  [UNDERLINE:text] → underlined text    ← likely the answer
  [BOLD:text]      → bold text          ← may be the answer (also used for headings)
  [ITALIC:text]    → italic text        ← may be the answer

Rules for formatted tags:
- Strip the tag wrapper: [RED:Malware] → correct_answer = "Malware"
- If multiple format tags appear inside one question, combine them: [RED:term1] [RED:term2] → "term1, term2"
- Ignore formatting on section headings (e.g. [BOLD:PART I. IDENTIFICATION]) — those are not answers
- Ignore formatting on item numbers (e.g. [BOLD:1.]) — those are not answers

Examples:
  "1. [RED:Photosynthesis] is the process by which plants make food."
  → question: "_______ is the process by which plants make food."  correct_answer: "Photosynthesis"

  "2. ________ [BOLD:osmosis] – the movement of water across a membrane"
  → correct_answer: "osmosis"

  "3. [HIGHLIGHT:CPU] stands for Central Processing Unit."
  → correct_answer: "CPU"

═══════════════════════════════════════════════════════
DASH / HYPHEN ANSWER PATTERNS
═══════════════════════════════════════════════════════
Teachers often embed the answer key using dashes. Detect these patterns:

Pattern A — answer BEFORE the dash:
  "Malware – This is malicious software that damages systems."
  "1. Malware - A type of malicious software."
  → question stem = the descriptive part, correct_answer = "Malware"

Pattern B — answer AFTER the blank/dash:
  "1. _______ – Malware – A type of malicious software."
  "1. _____ - Malware"
  → correct_answer = "Malware"  (the word/phrase between the blank and the description)

Pattern C — answer placed BEFORE the question number:
  "Malware – 1. What is a type of malicious software that damages systems?"
  → correct_answer = "Malware"

When extracting the question text for correct_answer detection:
- Remove the answer from the question stem and replace with a blank if needed
- The question field should contain the QUESTION TEXT only (not the answer inline)

═══════════════════════════════════════════════════════
QUESTION TYPES
═══════════════════════════════════════════════════════
Use ONLY these types: {', '.join(question_types)}

- "multiple_choice"  → has lettered options A B C D
- "true_false"       → asks true or false
- "identification"   → short answer; answer key may be formatted ([RED]/[BOLD]/[UNDERLINE]/[HIGHLIGHT]/[ITALIC]) or separated by a dash; also Scenario-Based and Spot-the-Error
- "essay"            → long open-ended answer (no single correct answer expected)
- "fill_blank"       → sentence with ___ to complete
- "matching"         → two-column table with Column A and Column B
- "enumeration"      → asks to LIST or ENUMERATE multiple items (e.g. "List the 9 types of...", "Give 5 examples of..."); correct_answer = newline-separated list of all items

═══════════════════════════════════════════════════════
MATCHING TYPE — CRITICAL INSTRUCTIONS
═══════════════════════════════════════════════════════
You will see matching tables formatted like this in the text:

    [TABLE]
    Column A  |  Column B
    1. CREATE  |  A. Permanently deletes a database object.
    2. TINYINT  |  B. Creates new database objects.
    3. DROP  |  C. Modifies the structure of an existing table.
    [/TABLE]

For ALL matching formats:
- Treat the ENTIRE section as ONE question with type "matching"
- Set "question" to the section heading (e.g. "Matching Type")
- Set "column_a" to ALL Column A items as a JSON array: ["1. CREATE", "2. TINYINT", ...]
- Set "column_b" to ALL Column B items as a JSON array: ["A. Permanently deletes...", "B. Creates...", ...]
- Set "matching_pairs" to [] if no answer key is present
- NEVER split a matching section into multiple identification questions

═══════════════════════════════════════════════════════
SCENARIO-BASED / SPOT THE ERROR / BONUS
═══════════════════════════════════════════════════════
Each numbered item = its own question.
type: "identification" or "essay"
Copy text exactly, including any code blocks.
correct_answer: "" unless explicitly shown.

═══════════════════════════════════════════════════════
TEXT TO SCAN:
═══════════════════════════════════════════════════════
{content}

═══════════════════════════════════════════════════════
RETURN — valid JSON only, no markdown, no extra text
═══════════════════════════════════════════════════════
{{
    "questions": [
        {{
            "type": "multiple_choice",
            "question": "exact question text",
            "options": {{"a": "...", "b": "...", "c": "...", "d": "..."}},
            "correct_answer": "a",
            "explanation": "",
            "difficulty": "medium",
            "points": 1
        }},
        {{
            "type": "matching",
            "question": "Matching Type",
            "column_a": ["1. CREATE", "2. TINYINT", "3. DROP", "4. DECIMAL(10,2)", "5. ALTER", "6. TRUNCATE", "7. CHAR(2)", "8. RENAME"],
            "column_b": [
                "A. Permanently deletes a database object.",
                "B. Creates new database objects.",
                "C. Modifies the structure of an existing table.",
                "D. Removes all rows but keeps table structure.",
                "E. Changes the name of a table.",
                "F. Storing exact currency values without rounding errors.",
                "G. Storing small whole numbers like age or status flags.",
                "H. Storing fixed-length codes such as country abbreviations (PH, US, UK)"
            ],
            "matching_pairs": [],
            "correct_answer": "",
            "explanation": "",
            "difficulty": "medium",
            "points": 1
        }},
        {{
            "type": "identification",
            "question": "_______ is the process by which plants make food using sunlight.",
            "correct_answer": "Photosynthesis",
            "explanation": "",
            "difficulty": "medium",
            "points": 1
        }},
        {{
            "type": "true_false",
            "question": "exact statement",
            "correct_answer": "true",
            "explanation": "",
            "difficulty": "easy",
            "points": 1
        }},
        {{
            "type": "fill_blank",
            "question": "The ___ command removes all rows but keeps table structure.",
            "correct_answer": "TRUNCATE",
            "explanation": "",
            "difficulty": "easy",
            "points": 1
        }},
        {{
            "type": "essay",
            "question": "exact essay question",
            "correct_answer": "",
            "explanation": "",
            "difficulty": "hard",
            "points": 5
        }},
        {{
            "type": "enumeration",
            "question": "List the nine (9) types of server security threats.",
            "correct_answer": "Malware\nDistributed Denial of Service (DDoS)\nUnauthorized Access\nSQL Injection\nPhishing\nInsider Threats\nZero-Day Exploits\nMisconfiguration\nBrute Force Attacks",
            "explanation": "",
            "difficulty": "medium",
            "points": 5
        }}
    ]
}}

JSON:"""

    def _build_generation_prompt(
        self,
        content: str,
        question_types: List[str],
        num_questions: int = 10,
    ) -> str:
        types_str = ', '.join(question_types)
        total     = num_questions * len(question_types)

        return f"""You are an expert teacher creating exam questions.

EDUCATIONAL CONTENT:
{content[:8000]}

TASK:
Generate exactly {num_questions} questions for EACH of these types: {types_str}
Total = {num_questions} × {len(question_types)} = {total} questions.

RULES:
- Base every question on the content above only.
- Vary difficulty: ~30% easy, ~50% medium, ~20% hard.
- multiple_choice : 4 options (a–d), one correct answer.
- true_false      : correct_answer = "true" or "false" (lowercase).
- identification  : correct_answer = specific term or phrase.
- essay           : correct_answer = key points / model answer.
- fill_blank      : use "___" in the question; correct_answer fills it.
- matching        : populate column_a, column_b, and matching_pairs.

Return ONLY valid JSON, no markdown.

{{
    "questions": [
        {{
            "type": "multiple_choice",
            "question": "...",
            "options": {{"a": "...", "b": "...", "c": "...", "d": "..."}},
            "correct_answer": "a",
            "explanation": "...",
            "difficulty": "medium",
            "points": 1
        }},
        {{
            "type": "matching",
            "question": "Match each item in Column A with the correct description in Column B.",
            "column_a": ["1. Term one", "2. Term two", "3. Term three"],
            "column_b": ["A. Description one", "B. Description two", "C. Description three"],
            "matching_pairs": [
                {{"item": "1. Term one",   "match": "A"}},
                {{"item": "2. Term two",   "match": "C"}},
                {{"item": "3. Term three", "match": "B"}}
            ],
            "correct_answer": "1-A, 2-C, 3-B",
            "explanation": "",
            "difficulty": "medium",
            "points": 1
        }}
    ]
}}

JSON:"""

    # =========================================================================
    # ORCHESTRATOR
    # =========================================================================

    def process_questionnaire(
        self,
        questionnaire,
        question_types: List[str],
        mode: str = 'extract',
        num_questions: int = 10,
    ):
        from questionnaires.models import ExtractedQuestion, QuestionType

        type_map = {}
        for qt in QuestionType.objects.filter(is_active=True):
            type_map[qt.name.lower()] = qt
            type_map[qt.name]         = qt

        aliases = {
            'multiple_choice':   'multiple_choice',
            'multiplechoice':    'multiple_choice',
            'mcq':               'multiple_choice',
            'true_false':        'true_false',
            'truefalse':         'true_false',
            'true/false':        'true_false',
            'fill_in_the_blank': 'fill_blank',
            'fill_blank':        'fill_blank',
            'fill in the blank': 'fill_blank',
            'identification':    'identification',
            'essay':             'essay',
            'matching':          'matching',
            'enumeration':       'enumeration',
            'short_answer':      'identification',
            'scenario':          'identification',
            'scenario-based':    'identification',
            'scenario_based':    'identification',
            'situation-based':   'identification',
            'spot the error':    'identification',
            'spot_the_error':    'identification',
            'bonus':             'identification',
            'problem solving':   'essay',
            'problem_solving':   'essay',
        }

        file_path = questionnaire.file.path
        content   = self.extract_text_from_file(file_path)

        if not content.strip():
            raise Exception("No text content could be extracted from the file")

        extracted_data = self.extract_questions_with_ai(
            content,
            question_types,
            mode=mode,
            num_questions=num_questions,
        )

        created_questions = []

        for q_data in extracted_data.get('questions', []):
            try:
                raw_type      = q_data.get('type', '').strip().lower()
                resolved_name = aliases.get(raw_type, raw_type)
                question_type = type_map.get(resolved_name) or type_map.get(raw_type)

                if not question_type:
                    print(f"Warning: type '{raw_type}' not in DB, skipping.")
                    continue

                option_a = option_b = option_c = option_d = None
                correct_answer = q_data.get('correct_answer', '')

                if resolved_name == 'matching':
                    col_a = q_data.get('column_a', [])
                    col_b = q_data.get('column_b', [])
                    pairs = q_data.get('matching_pairs', [])

                    option_a = json.dumps(col_a, ensure_ascii=False)
                    option_b = json.dumps(col_b, ensure_ascii=False)
                    option_c = json.dumps(pairs, ensure_ascii=False)
                    option_d = None

                    if not correct_answer and pairs:
                        correct_answer = ', '.join(
                            f"{p['item'].split('.')[0].strip()}-{p['match']}"
                            for p in pairs
                            if isinstance(p, dict) and 'item' in p and 'match' in p
                        )

                elif q_data.get('options'):
                    opts     = q_data['options']
                    option_a = opts.get('a')
                    option_b = opts.get('b')
                    option_c = opts.get('c')
                    option_d = opts.get('d')

                question = ExtractedQuestion.objects.create(
                    questionnaire  = questionnaire,
                    question_type  = question_type,
                    question_text  = q_data['question'],
                    option_a       = option_a,
                    option_b       = option_b,
                    option_c       = option_c,
                    option_d       = option_d,
                    correct_answer = correct_answer,
                    explanation    = q_data.get('explanation', ''),
                    difficulty     = q_data.get('difficulty', 'medium'),
                    points         = q_data.get('points', 1),
                )
                created_questions.append(question)

            except Exception as e:
                print(f"Error creating question: {e}")
                continue

        if not created_questions:
            raise Exception("No questions were found or created from the file")

        return created_questions