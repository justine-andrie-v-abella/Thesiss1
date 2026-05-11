# ============================================================================
# FILE: questionnaires/generators.py
# BISU Questionnaire Generator - Creates formatted DOCX and PDF files
# ============================================================================

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import os
import re
from django.conf import settings

TEMPLATE_PATH = r"C:\DJANGO PROJECTS\THESISattempt4\bisu_template.docx"

# ── Shared document formatting constants ─────────────────────────────────────
FONT_NAME = 'Arial Narrow'
FONT_SIZE = Pt(12)


class BISUQuestionnaireGenerator:

    def __init__(self, template_path=None):
        path = template_path or TEMPLATE_PATH
        if os.path.exists(path):
            self.doc = Document(path)
            print(f"✅ Loaded template: {path}")
        else:
            print(f"⚠️  Template not found at: {path}")
            self.doc = Document()
            self._setup_page_margins()

        # Apply global formatting defaults to the Normal style so every
        # paragraph and run inherits Arial Narrow 12 pt, single spacing, 0 pt
        # before/after — this covers both DOCX and the PDF rendered from it.
        self._setup_default_style()

    # =========================================================================
    # DOCUMENT-WIDE FORMATTING
    # =========================================================================

    def _setup_default_style(self):
        """Set Normal style to Arial Narrow 12pt, single line spacing, 0pt space before/after."""
        style = self.doc.styles['Normal']
        style.font.name = FONT_NAME
        style.font.size = FONT_SIZE
        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_before = Pt(0)
        pf.space_after  = Pt(0)

    @staticmethod
    def _fmt_para(p):
        """Apply single line spacing and 0pt before/after to a paragraph."""
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_before = Pt(0)
        pf.space_after  = Pt(0)

    @staticmethod
    def _fmt_run(r):
        """Apply Arial Narrow 12pt to a run."""
        r.font.name = FONT_NAME
        r.font.size = FONT_SIZE

    def _setup_page_margins(self):
        section = self.doc.sections[0]
        section.page_width    = Inches(8.5)
        section.page_height   = Inches(11)
        section.top_margin    = Inches(0.5)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    def generate_questionnaire(self, questionnaire_data):
        self._fill_placeholders(questionnaire_data)
        if 'sections' in questionnaire_data:
            answer_key_data = self._add_sections_in_order(questionnaire_data['sections'])
        else:
            answer_key_data = self._add_questions_by_type(questionnaire_data['questions'])
        if answer_key_data:
            self._add_answer_key_page(answer_key_data)
        return self.doc

    # =========================================================================
    # PLACEHOLDER REPLACEMENT
    # =========================================================================

    def _fill_placeholders(self, data):
        replacements = {
            '{{TITLE}}':       data.get('title', 'Examination'),
            '{{COURSE_CODE}}': data.get('course_code', ''),
            '{{COURSE_NAME}}': data.get('course_name', ''),
            '{{PROGRAM}}':     data.get('program', ''),
            '{{INSTRUCTOR}}':  data.get('instructor', ''),
            '{{DEPARTMENT}}':  data.get('department', ''),
            '{{SEMESTER}}':    data.get('semester', ''),
            '{{DIRECTIONS}}':  data.get('general_directions', self._get_default_directions()),
        }

        for paragraph in self.doc.paragraphs:
            self._replace_all(paragraph, replacements)

        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_all(paragraph, replacements)

        for section in self.doc.sections:
            for header_footer in [
                section.header, section.footer,
                section.even_page_header, section.even_page_footer,
                section.first_page_header, section.first_page_footer,
            ]:
                if header_footer is None:
                    continue
                for paragraph in header_footer.paragraphs:
                    self._replace_all(paragraph, replacements)
                for table in header_footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                self._replace_all(paragraph, replacements)

    def _replace_all(self, paragraph, replacements):
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                self._replace_in_paragraph(paragraph, placeholder, value)

    PLACEHOLDER_STYLES = {
        '{{TITLE}}': {
            'font_name': FONT_NAME,
            'font_size': FONT_SIZE,
            'bold':      True,
            'italic':    False,
            'underline': False,
        },
        '{{DEPARTMENT}}': {
            'font_name': FONT_NAME,
            'font_size': FONT_SIZE,
            'bold':      False,
            'italic':    False,
            'underline': False,
        },
    }

    def _replace_in_paragraph(self, paragraph, placeholder, value):
        full_text = ''.join(run.text for run in paragraph.runs)
        if placeholder not in full_text:
            return
        new_text = full_text.replace(placeholder, value)
        if paragraph.runs:
            first_run = paragraph.runs[0]
            first_run.text = new_text
            for run in paragraph.runs[1:]:
                run.text = ''
            style = self.PLACEHOLDER_STYLES.get(placeholder)
            if style:
                if 'font_name'  in style: first_run.font.name     = style['font_name']
                if 'font_size'  in style: first_run.font.size      = style['font_size']
                if 'bold'       in style: first_run.bold           = style['bold']
                if 'italic'     in style: first_run.italic         = style['italic']
                if 'underline'  in style: first_run.underline      = style['underline']
                if 'color'      in style: first_run.font.color.rgb = style['color']

    def _get_default_directions(self):
        return (
            "Write all your answers and solutions directly on the test questionnaire. "
            "Make sure your responses are well-organized and written clearly and legibly. "
            "After three (3) warnings, students caught discussing during the exam will be "
            "asked to IMMEDIATELY SURRENDER their test questionnaires. If you have any "
            "questions during the exam, feel free to ask the instructor for assistance. "
            "Wishing you all the best of luck on your exam!"
        )

    # =========================================================================
    # QUESTION SECTIONS
    # =========================================================================

    @staticmethod
    def _question_number_from_text(text):
        """Extract the leading question number/range from question_text."""
        m = re.match(r'^\s*(\d+\s*[\-–]\s*\d+|\d+)[\.\s]', text or '')
        return m.group(1).strip() if m else None

    def _add_sections_in_order(self, sections):
        """
        Render sections in the original document order.
        Each section has a 'header' (section_header text) and 'questions' list.
        Question numbers come from the question_text — the generator does NOT add its own.
        """
        answer_key = {}

        for idx, section in enumerate(sections):
            header_text = section.get('header') or ''
            questions   = section.get('questions', [])

            if not questions and not header_text:
                continue

            # ── Section header ──────────────────────────────────────────────
            if header_text:
                if '\n' in header_text:
                    title_line, _, instructions = header_text.partition('\n')
                    title_line   = title_line.strip()
                    instructions = instructions.strip()
                else:
                    m = re.match(
                        r'^((?:Part|PART)\s+[\w]+\.?\s+[^:]+?):\s*(.+)$',
                        header_text,
                        re.DOTALL,
                    )
                    if m:
                        title_line   = m.group(1).strip()
                        instructions = m.group(2).strip()
                    else:
                        title_line   = header_text
                        instructions = ''

                p = self.doc.add_paragraph()
                self._fmt_para(p)
                r = p.add_run(title_line)
                self._fmt_run(r)
                r.bold = True

                if instructions:
                    p2 = self.doc.add_paragraph()
                    self._fmt_para(p2)
                    r2 = p2.add_run(instructions)
                    self._fmt_run(r2)
                    r2.italic = True

                sp = self.doc.add_paragraph()
                self._fmt_para(sp)

            if not questions:
                continue

            section_answers = []
            for question in questions:
                self._add_question(question, number=None)
                qtype = question.question_type.name
                if qtype == 'essay':
                    continue
                qnum = self._question_number_from_text(question.question_text) or '?'
                if qtype == 'matching':
                    md    = question.get_matching_data()
                    pairs = md.get('pairs', []) if md else []
                    section_answers.append((qnum, 'matching', pairs))
                else:
                    section_answers.append((qnum, qtype, question.correct_answer or ''))

            if section_answers:
                if header_text:
                    if '\n' in header_text:
                        ak_title = header_text.partition('\n')[0].strip()
                    else:
                        m2 = re.match(r'^((?:Part|PART)\s+[\w]+\.?\s+[^:]+?):', header_text)
                        ak_title = m2.group(1).strip() if m2 else header_text
                else:
                    ak_title = f'Section {idx + 1}'
                answer_key[idx] = {
                    'title':   ak_title,
                    'answers': section_answers,
                }

            sp = self.doc.add_paragraph()
            self._fmt_para(sp)

        return answer_key

    def _add_questions_by_type(self, questions_by_type):
        """Adds all question sections and returns answer-key data grouped by type."""
        question_number = 1
        answer_key = {}

        for section_key, section_data in questions_by_type.items():
            if not section_data.get('questions'):
                continue

            # Section header (bold part title + instruction)
            p = self.doc.add_paragraph()
            self._fmt_para(p)
            r  = p.add_run(f"{section_data['title']}.")
            self._fmt_run(r)
            r.bold = True
            r2 = p.add_run(f" {section_data['instruction']}")
            self._fmt_run(r2)

            sp = self.doc.add_paragraph()
            self._fmt_para(sp)

            section_answers = []
            for question in section_data['questions']:
                self._add_question(question, question_number)
                qtype = question.question_type.name
                if qtype not in ('essay',):
                    if qtype == 'matching':
                        md = question.get_matching_data()
                        pairs = md.get('pairs', []) if md else []
                        section_answers.append((question_number, 'matching', pairs))
                    else:
                        section_answers.append(
                            (question_number, qtype, question.correct_answer or '')
                        )
                question_number += 1

            if section_answers:
                answer_key[section_key] = {
                    'title':   section_data['title'],
                    'answers': section_answers,
                }

            sp = self.doc.add_paragraph()
            self._fmt_para(sp)

        return answer_key

    def _add_question(self, question, number=None):
        """Dispatch to the correct renderer based on question type."""
        qtype = question.question_type.name

        if qtype == 'matching':
            self._add_matching_question(question, number)
            return

        # Build the stem text
        if number is not None:
            if qtype == 'true_false':
                stem = f"________ {number}. {question.question_text}"
            else:
                stem = f"{number}. {question.question_text}"
        else:
            if qtype == 'true_false':
                stem = f"________ {question.question_text}"
            else:
                stem = question.question_text

        p = self.doc.add_paragraph()
        self._fmt_para(p)
        r = p.add_run(stem)
        self._fmt_run(r)

        if qtype == 'multiple_choice':
            self._add_multiple_choice(question)

        elif qtype in ('identification', 'fill_blank', 'fill_in_the_blank'):
            p2 = self.doc.add_paragraph()
            self._fmt_para(p2)
            r2 = p2.add_run("Answer: ________________________")
            self._fmt_run(r2)

        elif qtype == 'enumeration':
            raw = question.correct_answer or ''
            items = [s.strip() for s in raw.splitlines() if s.strip()]
            if len(items) <= 1:
                items = [s.strip() for s in raw.split(',') if s.strip()]
            num_lines = max(len(items), 3)
            for idx in range(num_lines):
                p2 = self.doc.add_paragraph()
                self._fmt_para(p2)
                r2 = p2.add_run(f"{idx + 1}. ________________________")
                self._fmt_run(r2)

        elif qtype == 'essay':
            for _ in range(4):
                p2 = self.doc.add_paragraph()
                self._fmt_para(p2)
                r2 = p2.add_run("_" * 80)
                self._fmt_run(r2)

    # =========================================================================
    # MULTIPLE CHOICE
    # =========================================================================

    def _add_multiple_choice(self, question):
        opts = [
            ('a', question.option_a),
            ('b', question.option_b),
            ('c', question.option_c),
            ('d', question.option_d),
        ]
        table = self.doc.add_table(rows=2, cols=2)
        for i, (letter, text) in enumerate(opts):
            cell = table.rows[i // 2].cells[i % 2]
            p = cell.paragraphs[0]
            self._fmt_para(p)
            r = p.add_run(f"{letter}. {text or ''}")
            self._fmt_run(r)
        sp = self.doc.add_paragraph()
        self._fmt_para(sp)

    # =========================================================================
    # MATCHING TYPE
    # =========================================================================

    def _add_matching_question(self, question, start_number):
        md = question.get_matching_data()

        if not md:
            p = self.doc.add_paragraph()
            self._fmt_para(p)
            r = p.add_run(
                f"   [Could not render matching table. "
                f"Raw answer key: {question.correct_answer or '(none)'}]"
            )
            self._fmt_run(r)
            r.italic = True
            sp = self.doc.add_paragraph()
            self._fmt_para(sp)
            return

        column_a = md['column_a']
        column_b = md['column_b']
        pairs    = md['pairs']

        n_rows = max(len(column_a), len(column_b))
        if n_rows == 0:
            return

        section      = self.doc.sections[0]
        page_w       = section.page_width
        left_margin  = section.left_margin
        right_margin = section.right_margin
        content_w    = page_w - left_margin - right_margin
        content_w_dxa = int(content_w / 635)

        col_blank = int(content_w_dxa * 0.10)
        col_a     = int(content_w_dxa * 0.30)
        col_b     = content_w_dxa - col_blank - col_a

        table = self.doc.add_table(rows=1 + n_rows, cols=3)

        tbl    = table._tbl
        tblPr  = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'),    str(content_w_dxa))
        tblW.set(qn('w:type'), 'dxa')
        existing = tblPr.find(qn('w:tblW'))
        if existing is not None:
            tblPr.remove(existing)
        tblPr.append(tblW)

        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        existing = tblPr.find(qn('w:jc'))
        if existing is not None:
            tblPr.remove(existing)
        tblPr.append(jc)

        self._set_table_borders_invisible(tblPr)
        self._set_col_widths(table, [col_blank, col_a, col_b])

        # Header row
        header_labels = ['Ans.', 'Column A', 'Column B']
        for ci, label in enumerate(header_labels):
            cell = table.rows[0].cells[ci]
            p    = cell.paragraphs[0]
            self._fmt_para(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r    = p.add_run(label)
            self._fmt_run(r)
            r.bold      = True
            r.underline = True

        # Data rows
        for i in range(n_rows):
            row = table.rows[i + 1]

            cell_blank = row.cells[0]
            p = cell_blank.paragraphs[0]
            self._fmt_para(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("_______")
            self._fmt_run(r)

            a_text = column_a[i] if i < len(column_a) else ''
            cell_a = row.cells[1]
            p = cell_a.paragraphs[0]
            self._fmt_para(p)
            r = p.add_run(a_text)
            self._fmt_run(r)

            b_text = column_b[i] if i < len(column_b) else ''
            cell_b = row.cells[2]
            p = cell_b.paragraphs[0]
            self._fmt_para(p)
            r = p.add_run(b_text)
            self._fmt_run(r)

        sp = self.doc.add_paragraph()
        self._fmt_para(sp)

    # =========================================================================
    # TABLE HELPERS
    # =========================================================================

    @staticmethod
    def _set_col_widths(table, widths_dxa):
        """Set individual column widths (DXA) on every cell in each column."""
        for ci, width in enumerate(widths_dxa):
            for cell in table.columns[ci].cells:
                tc   = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW  = OxmlElement('w:tcW')
                tcW.set(qn('w:w'),    str(width))
                tcW.set(qn('w:type'), 'dxa')
                existing = tcPr.find(qn('w:tcW'))
                if existing is not None:
                    tcPr.remove(existing)
                tcPr.append(tcW)

    @staticmethod
    def _set_table_borders_invisible(tblPr):
        existing = tblPr.find(qn('w:tblBorders'))
        if existing is not None:
            tblPr.remove(existing)

        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'),   'none')
            border.set(qn('w:sz'),    '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tblBorders.append(border)

        tblPr.append(tblBorders)

    @staticmethod
    def _shade_cell(cell, fill_hex):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  fill_hex)
        existing = tcPr.find(qn('w:shd'))
        if existing is not None:
            tcPr.remove(existing)
        tcPr.append(shd)

    # =========================================================================
    # ANSWER KEY PAGE
    # =========================================================================

    def _add_answer_key_page(self, answer_key_data):
        # Hard page break
        p   = self.doc.add_paragraph()
        self._fmt_para(p)
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)

        # Page heading
        p = self.doc.add_paragraph()
        self._fmt_para(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('ANSWER KEY')
        self._fmt_run(r)
        r.bold = True

        p = self.doc.add_paragraph()
        self._fmt_para(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('(For Instructor Use Only — Detach Before Distributing)')
        self._fmt_run(r)
        r.italic         = True
        r.font.color.rgb = RGBColor(100, 100, 100)

        self._add_horizontal_rule()

        sp = self.doc.add_paragraph()
        self._fmt_para(sp)

        # Answers grouped by section
        for section_key, section_data in answer_key_data.items():
            answers = section_data.get('answers', [])
            if not answers:
                continue

            p = self.doc.add_paragraph()
            self._fmt_para(p)
            r = p.add_run(section_data['title'] + ':')
            self._fmt_run(r)
            r.bold = True

            matching_answers    = [(n, d) for n, qt, d in answers if qt == 'matching']
            enumeration_answers = [(n, d) for n, qt, d in answers if qt == 'enumeration']
            regular_answers     = [(n, d) for n, qt, d in answers
                                   if qt not in ('matching', 'enumeration')]

            if matching_answers:
                for num, pairs in matching_answers:
                    if not pairs:
                        continue
                    p = self.doc.add_paragraph()
                    self._fmt_para(p)
                    r = p.add_run(f"  {num}.")
                    self._fmt_run(r)
                    r.bold = True
                    parts = []
                    for pair in pairs:
                        item  = pair.get('item', '')
                        match = pair.get('match', '?')
                        n     = item.split('.')[0].strip() if '.' in str(item) else str(item)
                        parts.append(f"{n}→{match}")
                    p = self.doc.add_paragraph()
                    self._fmt_para(p)
                    r = p.add_run('   '.join(parts))
                    self._fmt_run(r)

            if enumeration_answers:
                for num, answer in enumeration_answers:
                    p = self.doc.add_paragraph()
                    self._fmt_para(p)
                    r = p.add_run(f"  {num}:")
                    self._fmt_run(r)
                    r.bold = True
                    raw_items = [s.strip() for s in answer.splitlines() if s.strip()]
                    if len(raw_items) <= 1:
                        raw_items = [s.strip() for s in answer.split(',') if s.strip()]
                    for i, item in enumerate(raw_items, 1):
                        p = self.doc.add_paragraph()
                        self._fmt_para(p)
                        r = p.add_run(f"     {i}. {item}")
                        self._fmt_run(r)

            if regular_answers:
                cols = 5
                rows_needed = (len(regular_answers) + cols - 1) // cols
                table = self.doc.add_table(rows=rows_needed, cols=cols)

                tbl   = table._tbl
                tblPr = tbl.find(qn('w:tblPr'))
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                    tbl.insert(0, tblPr)
                self._set_table_borders_invisible(tblPr)

                for idx, (num, answer) in enumerate(regular_answers):
                    row_i = idx // cols
                    col_i = idx % cols
                    cell  = table.rows[row_i].cells[col_i]
                    p     = cell.paragraphs[0]
                    self._fmt_para(p)
                    r     = p.add_run(f"{num}. {answer}")
                    self._fmt_run(r)

            sp = self.doc.add_paragraph()
            self._fmt_para(sp)

    def _add_horizontal_rule(self):
        p = self.doc.add_paragraph()
        self._fmt_para(p)
        r = p.add_run('─' * 55)
        self._fmt_run(r)
        r.font.color.rgb = RGBColor(150, 150, 150)

    # =========================================================================
    # SAVE METHODS
    # =========================================================================

    def save_to_buffer(self):
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer

    def save_docx(self, filepath):
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        self.doc.save(filepath)
        return filepath

    def save_pdf(self, filepath):
        docx_path = filepath.replace('.pdf', '_temp.docx')
        self.save_docx(docx_path)
        try:
            from docx2pdf import convert
            convert(docx_path, filepath)
            os.remove(docx_path)
            return filepath
        except ImportError:
            print("Warning: docx2pdf not installed. Returning DOCX instead.")
            return docx_path
        except Exception as e:
            print(f"PDF conversion failed: {e}")
            return docx_path


# =============================================================================
# MAIN ENTRY POINT
# =============================================================================

def generate_bisu_questionnaire(questionnaire_obj, selected_questions):
    """
    Generate BISU questionnaire from database objects using the Word template.
    Returns: tuple (docx_path, pdf_path)  — pdf_path may be None.
    """
    from .models import ExtractedQuestion

    selected_list = list(selected_questions)
    selected_ids  = {q.id for q in selected_list}

    section_headers = list(
        questionnaire_obj.extracted_questions
        .filter(question_type__name='section_header')
        .exclude(id__in=selected_ids)
        .select_related('question_type')
        .order_by('created_at')
    )

    if section_headers:
        all_questions = sorted(
            selected_list + section_headers,
            key=lambda q: q.created_at,
        )
    else:
        all_questions = selected_list

    sections = []
    current_header = None
    current_questions = []

    for q in all_questions:
        if q.question_type.name == 'section_header':
            if current_questions or current_header is not None:
                sections.append({
                    'header':    current_header,
                    'questions': current_questions[:],
                })
            current_header    = q.question_text
            current_questions = []
        else:
            current_questions.append(q)

    if current_questions or current_header is not None:
        sections.append({
            'header':    current_header,
            'questions': current_questions[:],
        })

    if not sections:
        sections = [{'header': None, 'questions': all_questions}]

    try:
        instructor_name = (questionnaire_obj.uploader.user.get_full_name()
                           or questionnaire_obj.uploader.user.username)
    except Exception:
        instructor_name = "Instructor"

    questionnaire_data = {
        'title':       questionnaire_obj.title or 'Examination',
        'course_code': questionnaire_obj.subject.code,
        'course_name': questionnaire_obj.subject.name,
        'program':     questionnaire_obj.department.code,
        'instructor':  instructor_name,
        'department':  questionnaire_obj.department.name,
        'semester':    '1st Semester, A.Y.2025-2026',
        'sections':    sections,
    }

    generator = BISUQuestionnaireGenerator()
    generator.generate_questionnaire(questionnaire_data)

    output_dir = os.path.join(settings.MEDIA_ROOT, 'generated_questionnaires')
    os.makedirs(output_dir, exist_ok=True)

    safe_name = f"{questionnaire_obj.subject.code}_{questionnaire_obj.title}".replace(' ', '_')
    safe_name = "".join(c for c in safe_name if c.isalnum() or c in ('_', '-'))

    docx_path = os.path.join(output_dir, f"{safe_name}.docx")
    pdf_path  = os.path.join(output_dir, f"{safe_name}.pdf")

    generator.save_docx(docx_path)

    try:
        generator.save_pdf(pdf_path)
    except Exception as e:
        print(f"PDF generation failed: {e}")
        pdf_path = None

    return docx_path, pdf_path
