# ============================================================================
# FILE: questionnaires/extractors.py
# AI-Powered Question Extraction System using Anthropic Claude API
# ============================================================================

import os
import json
import re
import logging
from typing import List, Dict, Any
from django.conf import settings
import anthropic

logger = logging.getLogger(__name__)

# File reading libraries
try:
    import PyPDF2
    from docx import Document
    import openpyxl
except ImportError:
    pass

try:
    import fitz  # PyMuPDF — optional, improves PDF colour/column detection
    _PYMUPDF_AVAILABLE = True
except ImportError:
    _PYMUPDF_AVAILABLE = False


class AIQuestionExtractor:
    """
    Scans uploaded files and extracts questions that are already written there.
    Does NOT generate or create new questions.
    Supports: PDF, DOCX, TXT, XLSX, XLS
    """

    CLAUDE_MODEL = 'claude-haiku-4-5-20251001'

    def __init__(self):
        api_key = getattr(settings, 'ANTHROPIC_API_KEY', None)
        if not api_key:
            raise ValueError(
                "ANTHROPIC_API_KEY not found in settings. "
                "Please add ANTHROPIC_API_KEY to your .env file."
            )
        self.client = anthropic.Anthropic(api_key=api_key)

    def process_questionnaire(self, questionnaire, type_names: List[str], mode: str = 'extract') -> List:
        """
        Process the questionnaire file.

        Args:
            questionnaire: Questionnaire model instance
            type_names: List of question type names to look for
            mode: 'extract' — copy questions already in the file (default)
                  'generate' — create new questions based on the file content

        Returns:
            List of created ExtractedQuestion objects
        """
        from questionnaires.models import ExtractedQuestion, QuestionType

        # Step 1: Read file content
        file_content = self._read_file(questionnaire.file.path, questionnaire.file_type)

        if not file_content.strip():
            raise ValueError("File is empty or could not be read")

        # Step 2: Extract or generate questions depending on mode
        extracted_data = self._extract_with_ai(file_content, type_names, mode=mode)

        # Step 3: Save questions to the database
        logger.info(
            "AI returned %d raw questions for questionnaire pk=%s",
            len(extracted_data),
            getattr(questionnaire, 'pk', '?'),
        )
        for i, qd in enumerate(extracted_data):
            logger.info(
                "  raw[%d]: type=%r question=%r col_a_len=%s col_b_len=%s",
                i,
                qd.get('type', '?'),
                str(qd.get('question', ''))[:60],
                len(qd.get('column_a', [])) if isinstance(qd.get('column_a'), list) else qd.get('column_a', 'N/A'),
                len(qd.get('column_b', [])) if isinstance(qd.get('column_b'), list) else qd.get('column_b', 'N/A'),
            )

        created_questions = []
        for question_data in extracted_data:
            try:
                q_type_name = question_data['type']
                question_type, created = QuestionType.objects.get_or_create(
                    name=q_type_name,
                    defaults={'description': '', 'is_active': True},
                )
                if created:
                    logger.warning(
                        "QuestionType '%s' did not exist — created it automatically",
                        q_type_name,
                    )

                # ── Matching type: store column_a/column_b/pairs as JSON ──────
                if q_type_name == 'matching':
                    col_a = question_data.get('column_a', [])
                    col_b = question_data.get('column_b', [])
                    pairs = question_data.get('matching_pairs', [])

                    import json as _json_inner
                    option_a = _json_inner.dumps(col_a, ensure_ascii=False) if col_a else None
                    option_b = _json_inner.dumps(col_b, ensure_ascii=False) if col_b else None
                    option_c = _json_inner.dumps(pairs, ensure_ascii=False)
                    option_d = None

                    correct_answer = question_data.get('answer', '')
                    # Build a compact answer-key string from pairs if AI didn't supply one
                    if not correct_answer and pairs:
                        correct_answer = ', '.join(
                            f"{p['item'].split('.')[0].strip()}-{p['match']}"
                            for p in pairs
                            if isinstance(p, dict) and 'item' in p and 'match' in p
                        )

                # ── All other types ────────────────────────────────────────────
                else:
                    option_a       = question_data.get('option_a')
                    option_b       = question_data.get('option_b')
                    option_c       = question_data.get('option_c')
                    option_d       = question_data.get('option_d')
                    correct_answer = question_data.get('answer', '')

                question = ExtractedQuestion.objects.create(
                    questionnaire=questionnaire,
                    question_type=question_type,
                    question_text=question_data['question'],
                    option_a=option_a,
                    option_b=option_b,
                    option_c=option_c,
                    option_d=option_d,
                    correct_answer=correct_answer,
                    explanation=question_data.get('explanation', ''),
                    points=question_data.get('points', 1),
                    difficulty=question_data.get('difficulty', 'medium'),
                    is_approved=False
                )
                created_questions.append(question)
            except Exception as e:
                logger.error("Error saving question: %s", e, exc_info=True)
                continue

        return created_questions

    def _read_file(self, file_path: str, file_type: str) -> str:
        """Read content from various file types"""
        if file_type == 'txt':
            return self._read_txt(file_path)
        elif file_type == 'pdf':
            return self._read_pdf(file_path)
        elif file_type in ['docx', 'doc']:
            return self._read_docx(file_path)
        elif file_type in ['xlsx', 'xls']:
            return self._read_xlsx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    # -------------------------------------------------------------------------
    # Formatting detection helpers (used by both PDF and DOCX readers)
    # -------------------------------------------------------------------------

    @staticmethod
    def _get_run_format_tag(run):
        """
        Inspect a python-docx Run for answer-key formatting cues.
        Returns 'RED', 'HIGHLIGHT', 'UNDERLINE', 'BOLD', or 'ITALIC' — or None
        if the run has no special formatting.

        Priority: RED > HIGHLIGHT > UNDERLINE > BOLD > ITALIC
        (RED colour is the strongest signal that something is an answer key.)
        """
        try:
            from docx.oxml.ns import qn
            rPr = run._r.find(qn('w:rPr'))
            if rPr is not None:

                # 1. Direct RGB colour ------------------------------------------
                color_el = rPr.find(qn('w:color'))
                if color_el is not None:
                    val = (color_el.get(qn('w:val')) or '').strip()
                    if val and val.lower() not in ('auto', '000000', 'ffffff') and len(val) == 6:
                        try:
                            r_v = int(val[0:2], 16)
                            g_v = int(val[2:4], 16)
                            b_v = int(val[4:6], 16)
                            if r_v > 150 and g_v < 100 and b_v < 100:
                                return 'RED'
                        except ValueError:
                            pass

                # 2. Highlight --------------------------------------------------
                hl = rPr.find(qn('w:highlight'))
                if hl is not None:
                    hl_val = (hl.get(qn('w:val')) or '').lower()
                    if hl_val and hl_val != 'none':
                        return 'HIGHLIGHT'

        except Exception:
            pass

        # 3. Underline (python-docx API is reliable here) ----------------------
        try:
            if run.underline:
                return 'UNDERLINE'
        except Exception:
            pass

        # 4. Bold (lower priority — headings are also bold) --------------------
        try:
            if run.bold:
                return 'BOLD'
        except Exception:
            pass

        # 5. Italic ------------------------------------------------------------
        try:
            if run.italic:
                return 'ITALIC'
        except Exception:
            pass

        return None

    @staticmethod
    def _is_noise_run(text: str) -> bool:
        """
        Returns True for short runs that are unlikely to be answer text even if
        formatted — e.g. question numbers ("1.", "2)"), pure punctuation, blanks.
        These are skipped to avoid tagging noise as answers.
        """
        stripped = text.strip()
        if not stripped:
            return True
        # Purely numeric / punctuation runs (question numbers, dashes, underscores)
        if re.match(r'^[\d\s\.\)\-\_]+$', stripped):
            return True
        # Very short single non-alpha character
        if len(stripped) <= 1 and not stripped.isalpha():
            return True
        return False

    def _extract_para_with_formatting(self, para) -> str:
        """
        Extract paragraph text, wrapping formatted runs in tags so Claude AI
        can identify them as answer-key markers.
        e.g.  "1. [RED:Malware] is a type of malicious software."
        """
        parts = []
        for run in para.runs:
            text = run.text
            if not text:
                continue
            tag = self._get_run_format_tag(run)
            if tag and not self._is_noise_run(text):
                stripped = text.strip()
                leading  = text[: len(text) - len(text.lstrip())]
                trailing = text[len(text.rstrip()):]
                parts.append(f"{leading}[{tag}:{stripped}]{trailing}")
            else:
                parts.append(text)
        result = ''.join(parts).strip()
        # Fallback: use para.text if all runs were plain (handles edge cases)
        return result if result else para.text.strip()

    def _read_txt(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    def _read_pdf(self, file_path: str) -> str:
        """
        Extract text from a PDF with formatting cues preserved where possible.

        Uses PyMuPDF (fitz) as the primary engine because it:
          - Detects text colour → emits [RED:text] tags for reddish spans
          - Handles multi-column layouts correctly (sorts by position)
        Falls back to PyPDF2 if PyMuPDF is unavailable.
        """
        if _PYMUPDF_AVAILABLE:
            try:
                return self._read_pdf_pymupdf(file_path)
            except Exception as e:
                logger.warning("PyMuPDF failed (%s), falling back to PyPDF2", e)

        # ── PyPDF2 fallback ────────────────────────────────────────────────
        text = []
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
        except Exception as e:
            raise ValueError(f"Failed to read PDF: {str(e)}")
        return '\n\n'.join(text)

    def _read_pdf_pymupdf(self, file_path: str) -> str:
        """
        PyMuPDF-based PDF extraction.

        Key features:
          1. Colour detection — spans with reddish colour get tagged [RED:text]
          2. Multi-column layout — text blocks are sorted by their vertical
             position within each horizontal column so columns read top-to-bottom
             rather than being interleaved.
          3. Bold / italic detection via font flags.
        """
        COLUMN_BAND_TOLERANCE = 15  # px — spans within this y-delta share a row

        def is_reddish(color_int):
            """Return True for colours where R > 150, G < 100, B < 100."""
            if color_int is None:
                return False
            r = (color_int >> 16) & 0xFF
            g = (color_int >> 8)  & 0xFF
            b =  color_int        & 0xFF
            return r > 150 and g < 100 and b < 100

        def span_tag(span):
            """Return formatting tag for a span, or None."""
            color = span.get('color', 0)
            if is_reddish(color):
                return 'RED'
            flags = span.get('flags', 0)
            # fitz font flags: bit 4 = italic, bits 3/1 = bold
            if flags & (1 << 4):
                return 'ITALIC'
            if flags & (1 << 3) or flags & (1 << 1):
                return 'BOLD'
            return None

        def spans_to_lines(spans):
            """Cluster spans into logical lines by y0 proximity."""
            if not spans:
                return []
            spans = sorted(spans, key=lambda s: (round(s[1] / 5) * 5, s[0]))
            lines    = []
            cur_y    = None
            cur_line = []
            for span in spans:
                y = span[1]
                if cur_y is None or abs(y - cur_y) > COLUMN_BAND_TOLERANCE:
                    if cur_line:
                        lines.append(cur_line)
                    cur_line = [span]
                    cur_y    = y
                else:
                    cur_line.append(span)
            if cur_line:
                lines.append(cur_line)
            return lines

        def render_lines(lines):
            result = []
            for line in lines:
                line_text = ''
                for (x0, y0, x1, y1, txt, tag) in sorted(line, key=lambda s: s[0]):
                    stripped = txt.strip()
                    if not stripped:
                        line_text += txt
                        continue
                    if tag and not self._is_noise_run(stripped):
                        leading  = txt[: len(txt) - len(txt.lstrip())]
                        trailing = txt[len(txt.rstrip()):]
                        line_text += f'{leading}[{tag}:{stripped}]{trailing}'
                    else:
                        line_text += txt
                rendered = line_text.strip()
                if rendered:
                    result.append(rendered)
            return result

        pages_text = []
        try:
            doc = fitz.open(file_path)
            for page in doc:
                page_width = page.rect.width

                # Collect all text spans with position info ──────────────────
                raw_spans = []  # (x0, y0, x1, y1, text, tag)
                blocks = page.get_text('dict', flags=fitz.TEXT_PRESERVE_WHITESPACE)['blocks']
                for block in blocks:
                    if block.get('type') != 0:
                        continue
                    for line in block.get('lines', []):
                        for span in line.get('spans', []):
                            txt = span.get('text', '')
                            if not txt.strip():
                                continue
                            bbox = span.get('bbox', (0, 0, 0, 0))
                            tag  = span_tag(span)
                            raw_spans.append((bbox[0], bbox[1], bbox[2], bbox[3], txt, tag))

                if not raw_spans:
                    continue

                mid_x       = page_width / 2
                left_spans  = [s for s in raw_spans if s[0] < mid_x]
                right_spans = [s for s in raw_spans if s[0] >= mid_x]

                # If there are meaningful right-column spans, render columns separately
                if right_spans and len(right_spans) >= 2:
                    left_lines  = spans_to_lines(left_spans)
                    right_lines = spans_to_lines(right_spans)
                    page_lines  = render_lines(left_lines) + render_lines(right_lines)
                else:
                    all_lines  = spans_to_lines(raw_spans)
                    page_lines = render_lines(all_lines)

                pages_text.append('\n'.join(page_lines))

            doc.close()
            return '\n\n'.join(pages_text)

        except Exception as e:
            raise ValueError(f"PyMuPDF error: {str(e)}") from e

    def _read_docx(self, file_path: str) -> str:
        """
        Extract DOCX text in reading order, preserving answer-key formatting
        cues ([RED:...], [BOLD:...], etc.).
        Also reads table cells so answer keys placed in tables are captured.
        """
        _WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

        def _plain_text(element):
            """Fallback: extract plain text from an XML element."""
            return ''.join(
                node.text for node in element.iter()
                if node.tag.endswith('}t') and node.text
            ).strip()

        try:
            doc = Document(file_path)
            text = []

            # Build element → paragraph object map so we can use
            # _extract_para_with_formatting (which needs a python-docx Para).
            # doc.paragraphs includes ALL paragraphs, including those in tables.
            para_map = {p._element: p for p in doc.paragraphs}

            for block in doc.element.body:
                tag = block.tag.split('}')[-1] if '}' in block.tag else block.tag

                if tag == 'p':
                    para = para_map.get(block)
                    if para is not None:
                        para_text = self._extract_para_with_formatting(para)
                    else:
                        para_text = _plain_text(block)

                    if para_text:
                        text.append(para_text)

                elif tag == 'tbl':
                    # Table — read every cell row by row with formatting cues
                    for row in block.iter(f'{{{_WNS}}}tr'):
                        row_cells = []
                        for cell in row.iter(f'{{{_WNS}}}tc'):
                            cell_parts = []
                            for para_elem in cell.iter(f'{{{_WNS}}}p'):
                                para = para_map.get(para_elem)
                                if para is not None:
                                    pt = self._extract_para_with_formatting(para)
                                else:
                                    pt = _plain_text(para_elem)
                                if pt:
                                    cell_parts.append(pt)
                            cell_text = ' '.join(cell_parts).strip()
                            if cell_text:
                                row_cells.append(cell_text)
                        if row_cells:
                            text.append('  |  '.join(row_cells))

            return '\n'.join(text)
        except Exception as e:
            raise ValueError(f"Failed to read DOCX: {str(e)}")

    def _read_xlsx(self, file_path: str) -> str:
        try:
            workbook = openpyxl.load_workbook(file_path)
            text = []
            for sheet in workbook.worksheets:
                text.append(f"\n=== {sheet.title} ===\n")
                for row in sheet.iter_rows(values_only=True):
                    row_text = '\t'.join([str(cell) if cell is not None else '' for cell in row])
                    if row_text.strip():
                        text.append(row_text)
            return '\n'.join(text)
        except Exception as e:
            raise ValueError(f"Failed to read XLSX: {str(e)}")

    def _extract_with_ai(self, content: str, type_names: List[str], mode: str = 'extract') -> List[Dict[str, Any]]:
        """
        Scan or generate questions depending on mode.
        mode='extract'  → copies questions already written in the file
        mode='generate' → creates new questions based on the file content
        """

        # Claude Haiku supports up to 200k tokens.
        # Always preserve the END of the document — that's where answer keys live.
        MAX_CHARS = 150_000
        if len(content) > MAX_CHARS:
            head = content[:100_000]
            tail = content[-50_000:]
            content = head + "\n\n... (middle section truncated) ...\n\n" + tail

        if mode == 'generate':
            prompt = self._build_generation_prompt(content, type_names)
            temperature = 0.7
        else:
            prompt = self._build_extraction_prompt(content, type_names)
            temperature = 0.1

        try:
            response = self.client.messages.create(
                model=self.CLAUDE_MODEL,
                max_tokens=8192,
                temperature=temperature,
                system=(
                    "You are a document scanner that copies text exactly as written. "
                    "You extract questions and copy answers verbatim from the answer key — "
                    "you do NOT judge, evaluate, or skip any answer no matter what it says. "
                    "You always respond with valid JSON only — no explanation, no markdown."
                ),
                messages=[
                    {"role": "user",      "content": prompt},
                    {"role": "assistant", "content": "["},   # force JSON array start
                ],
            )

            # Claude prefill: the response continues after our "[" prefix
            raw = "[" + response.content[0].text
            questions = self._parse_ai_response(raw)
            return questions

        except Exception as e:
            logger.error("Anthropic API error: %s", e, exc_info=True)
            raise Exception(f"AI extraction failed: {str(e)}") from e

    def _build_extraction_prompt(self, content: str, type_names: List[str]) -> str:
        """
        Prompt that tells Claude to copy questions from the file and match
        answers from any answer key section (inline or end-of-document).
        """

        type_descriptions = {
            'multiple_choice': 'Multiple Choice (has options A, B, C, D)',
            'true_false':      'True/False (answer is True or False)',
            'identification':  'Identification (short answer, one word or phrase); answer key may be formatted ([RED]/[BOLD]/[UNDERLINE]/[HIGHLIGHT]/[ITALIC]) or separated by a dash',
            'essay':           'Essay (requires a long written answer)',
            'fill_blank':      'Fill in the Blank (sentence with a blank to complete)',
            'matching':        'Matching Type (match items from two columns)',
            'enumeration':     'Enumeration (asks to LIST or ENUMERATE multiple items, e.g. "List the 9 types of..."); answer = newline-separated list of all items',
        }

        types_list = '\n'.join([
            f"- {type_descriptions.get(t, t)}" for t in type_names
        ])

        prompt = f"""You are scanning a test questionnaire document. Your job is to:
  1. Extract every question exactly as written.
  2. Find the answer for each question from the answer key — either a separate key section
     OR inline formatting/dash patterns within the question itself.

QUESTION TYPES TO FIND:
{types_list}

════════════════════════════════════════
STEP 1 — LOCATE THE ANSWER KEY FIRST
════════════════════════════════════════
Before extracting questions, scan the ENTIRE document for an answer key section.
It can appear anywhere — beginning, middle, or end — and may be labeled:
  "Answer Key", "ANSWER KEY", "Answers", "Key", "Answer Sheet",
  or just a numbered/lettered list after the questions.

For each section (e.g. Part I, Part II, Section A, Section B), record the answers
IN ORDER with their numbers so you can match them to questions by position.

Example answer key formats you might see:
  • "1. Subset   2. True   3. Cartesian Product"
  • "Part I: 1. Haha   2. Hehe   3. Something"
  • "Part II: 1-A  2-C  3-B  4-D"
  • A plain list: "Haha / Hehe / Shader" (match by position: 1st item = question 1)

════════════════════════════════════════
STEP 2 — EXTRACT QUESTIONS
════════════════════════════════════════
- Copy every question text as written — do NOT rephrase or invent anything.
- Do NOT treat answer key entries as standalone questions.
- When an answer is embedded inline (via formatting tags or a dash pattern),
  extract that answer into the "answer" field and write the question WITHOUT
  the inline answer — replace it with a blank _______ if needed.
- If no questions are found at all, return [].

NUMBERING — PRESERVE THE ORIGINAL:
  ▸ Keep the EXACT question number that appears in the document in the "question" field.
    e.g. if the document says "15. Which of the following..." → question starts with "15."
  ▸ Do NOT renumber, do NOT restart from 1, do NOT add your own numbering.
  ▸ If the document has no number on a question, leave it without one.

PART / SECTION TRACKING — EXTRACT HEADER ONCE, NOT ON EVERY QUESTION:
  ▸ Detect section headers such as: "Part I", "Part II", "PART I.", "Section A",
    "I. Multiple Choice", "II. Identification", etc.
  ▸ When you reach a new section, emit ONE "section_header" entry BEFORE the questions in
    that section. Include the section title AND the directions/instructions line if present.
      {{"type": "section_header", "question": "Part I. Multiple Choice\nDirections: Choose the letter of the best answer.", "answer": "", "explanation": "", "difficulty": "easy", "points": 0}}
  ▸ Do NOT add the section label to individual question texts — questions start with their number only:
      "1. Which of the following best describes a file server?"   ← correct
      "(Part I) 1. Which of the following..."                     ← WRONG, do not do this
  ▸ If no section header exists in the document, do not emit any section_header entry.

⚠ CRITICAL — MATCHING TYPE RULE (DO NOT SKIP):
  If you see ANY of these signals, you are looking at a MATCHING TYPE section.
  You MUST extract it as ONE question with type "matching" — NEVER as individual questions:
    ▸ Headers "Column A" and "Column B" anywhere in the document
    ▸ Items formatted as "_X_ N. Term" (e.g. "_C_ 1. Linux Server", "_A_ 2. Windows Server")
      where X is a letter and N is a number — this IS a matching section with embedded answers
    ▸ A scoring/range note like "30 – 31. 5 correct = 2 pts | 3-4 correct = 1.5 pts..."
      followed by Column A / Column B lists — "30 – 31" is the question NUMBER RANGE
    ▸ A visible table with two columns of terms and descriptions

  DO NOT extract individual rows of a matching table as identification questions.
  DO NOT skip a matching section just because the columns appear on separate lines.
  DO process EACH matching section as its own separate "matching" question.

════════════════════════════════════════
FORMATTED TEXT — ANSWER KEY DETECTION
════════════════════════════════════════
The extracted text may contain formatting tags where the teacher used special
formatting to mark the answer. Use these to set the "answer" field:

  [RED:text]       → red-coloured text  ← STRONGEST signal; almost always the answer
  [HIGHLIGHT:text] → highlighted text   ← very strong signal; likely the answer
  [UNDERLINE:text] → underlined text    ← likely the answer
  [BOLD:text]      → bold text          ← may be the answer (also used for headings)
  [ITALIC:text]    → italic text        ← may be the answer

Rules for formatted tags:
- Strip the tag wrapper: [RED:Malware] → answer = "Malware"
- If multiple format tags appear inside one question, combine them:
    [RED:term1] [RED:term2] → "term1, term2"
- Ignore formatting on section headings (e.g. [BOLD:PART I. IDENTIFICATION])
- Ignore formatting on item numbers (e.g. [BOLD:1.])
- For the "question" field, remove the tag and replace with _______ where it was inline

Examples:
  "1. [RED:Photosynthesis] is the process by which plants make food."
  → question: "1. _______ is the process by which plants make food."
    answer: "Photosynthesis"

  "2. ________ [BOLD:osmosis] – the movement of water across a membrane"
  → answer: "osmosis"

  "3. [HIGHLIGHT:CPU] stands for Central Processing Unit."
  → answer: "CPU"

════════════════════════════════════════
MATCHING TYPE — CRITICAL INSTRUCTIONS
════════════════════════════════════════
A matching question has two columns of items that students must pair together.
Treat the ENTIRE matching section as ONE question — do NOT split each row into a
separate question.

Teachers may format matching sections in two ways:
  A) As a Word/PDF table with two columns (Column A | Column B)
  B) As two separate paragraph lists (all Column A items, then all Column B items)
  C) As text with the format:  "1. Term   |   A. Description"

For matching type, include these EXTRA fields in the JSON object:
  "column_a"       → array of ALL Column A items (terms/words to identify), WITHOUT any
                     embedded answer prefix — e.g. ["1. A Record", "2. MX Record", "3. DORA"]
  "column_b"       → array of ALL Column B items (descriptions to match)
                     e.g. ["A. The four-step process...", "B. A command...", "C. The specific DNS record..."]
  "matching_pairs" → array of answer pairs, built from wherever the answers appear:
                     [{{"item": "1. A Record", "match": "C"}}, {{"item": "2. MX Record", "match": "E"}}]
                     Set to [] if truly no answer key is found anywhere.

════════════════════════════════════════
MATCHING — ANSWER KEY FORMATS TO DETECT
════════════════════════════════════════

Format 1 — Underscore-letter prefix embedded in Column A (MOST COMMON in these files):
  "_C_ 1. A Record"   → column_a item = "1. A Record",   pair = item "1. A Record" → match "C"
  "_E_ 2. MX Record"  → column_a item = "2. MX Record",  pair = item "2. MX Record" → match "E"
  "_A_ 3. DORA"       → column_a item = "3. DORA",        pair = item "3. DORA"       → match "A"
  The letter between the underscores IS the answer — strip the "_X_" prefix when building column_a.

Format 2 — Separate answer key section (e.g. at end of document):
  "1-C  2-E  3-A  4-B  5-D"
  → matching_pairs = [{{"item":"1","match":"C"}}, {{"item":"2","match":"E"}}, ...]

Format 3 — In-table answer column (3-column layout: Answer | Column A | Column B):
  "C  |  1. A Record  |  A. The four-step process..."
  → extract answer letter "C" for item "1. A Record"

Rules:
  - Strip the "_X_" prefix from column_a items — do NOT include it in the item text
  - Preserve the numbering/lettering exactly (e.g. "1. Linux Server" stays "1. Linux Server")
  - The "question" field = section prefix (if any) + the question number range + scoring note:
      "(Part IV) 30-31. 5 correct = 2 pts | 3-4 correct = 1.5 pts | 1-2 correct = 1 pt | 0 correct = 0 pts"
    If there is no scoring note, use: "(Part IV) 30-31. Matching Type"
    If there is no range number, use the section heading alone.
  - The "answer" field = compact key string built from pairs, e.g. "1-C, 2-A, 3-B, 4-E, 5-D"
    Leave "" only if truly no answer information is found anywhere.
  - NEVER return individual matching rows as separate identification questions
  - A document may have MULTIPLE separate matching sections — treat each as its own question

Example document text:
  30 – 31. 5 correct = 2 pts | 3-4 correct = 1.5 pts | 1-2 correct = 1 pt | 0 correct = 0 pts
  Column A                      Column B
  _C_ 1. Linux Server           A. Proprietary OS; best for .NET and Active Directory integration.
  _A_ 2. Windows Server         B. Highly stable, proprietary OS often used in high-end workstations.
  _B_ 3. UNIX                   C. Open-source; favored for web servers due to low cost and security.
  _E_ 4. Command Line (CLI)     D. A minimal Windows install option that reduces the attack surface.
  _D_ 5. Server Core            E. The primary management interface for Linux/UNIX servers.

Expected JSON output for that section:
  {{
    "type": "matching",
    "question": "30-31. 5 correct = 2 pts | 3-4 correct = 1.5 pts | 1-2 correct = 1 pt | 0 correct = 0 pts",
    "column_a": ["1. Linux Server", "2. Windows Server", "3. UNIX", "4. Command Line (CLI)", "5. Server Core"],
    "column_b": ["A. Proprietary OS; best for .NET and Active Directory integration.", "B. Highly stable, proprietary OS often used in high-end workstations.", "C. Open-source; favored for web servers due to low cost and security.", "D. A minimal Windows install option that reduces the attack surface.", "E. The primary management interface for Linux/UNIX servers."],
    "matching_pairs": [{{"item": "1. Linux Server", "match": "C"}}, {{"item": "2. Windows Server", "match": "A"}}, {{"item": "3. UNIX", "match": "B"}}, {{"item": "4. Command Line (CLI)", "match": "E"}}, {{"item": "5. Server Core", "match": "D"}}],
    "answer": "1-C, 2-A, 3-B, 4-E, 5-D",
    "explanation": "",
    "difficulty": "medium",
    "points": 2
  }}

════════════════════════════════════════
DASH / HYPHEN ANSWER PATTERNS
════════════════════════════════════════
Teachers often embed the answer key using dashes. Detect these patterns:

Pattern A — answer BEFORE the dash (most common for identification):
  "Malware – This is malicious software that damages systems."
  "1. Malware - A type of malicious software."
  → question = the descriptive part, answer = "Malware"

Pattern B — answer AFTER the blank/dash:
  "1. _______ – Malware – A type of malicious software."
  "1. _____ - Malware"
  → answer = "Malware"  (the word/phrase right after the blank)

Pattern C — answer placed BEFORE the question number:
  "Malware – 1. What is a type of malicious software that damages systems?"
  → answer = "Malware"

When using a dash pattern, the "question" field should contain the question text
only (not the answer inline). Replace the answer position with _______ if needed.

════════════════════════════════════════
STEP 3 — COPY ANSWERS VERBATIM
════════════════════════════════════════
CRITICAL RULE: Copy whatever text the answer key says for each question.
Do NOT evaluate whether the answer is correct. Do NOT skip an answer because it
seems wrong, unusual, or doesn't match the question topic.
The answer key is the authority — copy it exactly, always.

Matching rules:
  • Answer keys typically use the ORIGINAL document question numbers (not restarted).
    Match each answer to the question with the SAME number in the document.
    e.g. answer key "15. A" → the question numbered "15." in the document
  • If the answer key DOES restart numbering per section (e.g. "Part I: 1. A  2. B..."),
    match by position within that section.
  • If the answer key has no numbers (plain list), match by position:
      1st answer item → 1st question in that section
      2nd answer item → 2nd question in that section, etc.
  • Multiple choice → store only the letter: "A", "B", "C", or "D" (uppercase)
  • Enumeration → store all items as a newline-separated list in the "answer" field
  • All other types → copy the answer text as written in the answer key
  • If truly no answer key entry exists for a question anywhere, use ""

TEXT TO SCAN:
{content}

════════════════════════════════════════
OUTPUT FORMAT
════════════════════════════════════════
Return ONLY a valid JSON array — no extra text, no markdown, no code fences.

Common fields for ALL question types:
  "type"        - one of {type_names} OR "section_header" (for part/section title + directions)
  "question"    - question text (without inline answer; use _______ placeholder)
                  for section_header: the section title + newline + directions line
  "answer"      - answer extracted from key, formatting tag, or dash pattern; or ""
  "explanation" - any explanation text found in the document, or ""
  "difficulty"  - "easy", "medium", or "hard"
  "points"      - point value if shown, otherwise 1 (use 0 for section_header)

Type-specific extra fields:
  multiple_choice only → "option_a", "option_b", "option_c", "option_d"
  matching only        → "column_a" (array), "column_b" (array), "matching_pairs" (array)
  section_header       → no extra fields needed

Examples (section header appears once; individual questions keep only their original number):
[
  {{
    "type": "section_header",
    "question": "Part I. Multiple Choice\nDirections: Choose the letter of the best answer. Write the letter on the blank provided.",
    "answer": "",
    "explanation": "",
    "difficulty": "easy",
    "points": 0
  }},
  {{
    "type": "multiple_choice",
    "question": "1. Which of the following is a type of malware?",
    "option_a": "Firewall",
    "option_b": "Virus",
    "option_c": "Router",
    "option_d": "Switch",
    "answer": "B",
    "explanation": "",
    "difficulty": "easy",
    "points": 1
  }},
  {{
    "type": "section_header",
    "question": "Part II. Identification\nDirections: Write the correct answer on the blank.",
    "answer": "",
    "explanation": "",
    "difficulty": "easy",
    "points": 0
  }},
  {{
    "type": "identification",
    "question": "16. _______ is the process by which plants make food using sunlight.",
    "answer": "Photosynthesis",
    "explanation": "",
    "difficulty": "medium",
    "points": 1
  }},
  {{
    "type": "section_header",
    "question": "Part IV. Matching Type\nDirections: Match Column A with Column B.",
    "answer": "",
    "explanation": "",
    "difficulty": "easy",
    "points": 0
  }},
  {{
    "type": "matching",
    "question": "30-31. 5 correct = 2 pts | 3-4 correct = 1.5 pts | 1-2 correct = 1 pt | 0 correct = 0 pts",
    "column_a": ["1. Linux Server", "2. Windows Server", "3. UNIX", "4. Command Line (CLI)", "5. Server Core"],
    "column_b": ["A. Proprietary OS; best for .NET and Active Directory integration.", "B. Highly stable, proprietary OS often used in high-end workstations.", "C. Open-source; favored for web servers due to low cost and security.", "D. A minimal Windows install option that reduces the attack surface.", "E. The primary management interface for Linux/UNIX servers."],
    "matching_pairs": [{{"item": "1. Linux Server", "match": "C"}}, {{"item": "2. Windows Server", "match": "A"}}, {{"item": "3. UNIX", "match": "B"}}, {{"item": "4. Command Line (CLI)", "match": "E"}}, {{"item": "5. Server Core", "match": "D"}}],
    "answer": "1-C, 2-A, 3-B, 4-E, 5-D",
    "explanation": "",
    "difficulty": "medium",
    "points": 2
  }}
]

If no questions found, return: []"""

        return prompt

    def _build_generation_prompt(self, content: str, type_names: List[str]) -> str:
        """
        Prompt for GENERATING new questions based on the file content.
        Used by the generate_questionnaire view.
        """

        type_descriptions = {
            'multiple_choice': 'Multiple Choice (4 options A, B, C, D)',
            'true_false':      'True/False',
            'identification':  'Identification (short answer)',
            'essay':           'Essay (detailed answer)',
            'fill_blank':      'Fill in the Blank',
            'matching':        'Matching Type',
        }

        types_list = '\n'.join([
            f"- {type_descriptions.get(t, t)}" for t in type_names
        ])

        prompt = f"""You are an expert teacher. Based on the educational content below, generate high-quality exam questions.

QUESTION TYPES TO GENERATE:
{types_list}

CONTENT TO BASE QUESTIONS ON:
{content}

INSTRUCTIONS:
1. Generate questions that test understanding of the content
2. For each question provide:
   - "type": one of {type_names}
   - "question": the question text
   - "option_a", "option_b", "option_c", "option_d": all 4 options (multiple choice only)
   - "answer": the correct answer
   - "explanation": brief explanation of why the answer is correct
   - "difficulty": "easy", "medium", or "hard"
   - "points": 1 for easy, 2 for medium, 3 for hard

3. Generate a good mix of difficulties
4. Return ONLY a valid JSON array, no extra text

JSON:"""

        return prompt

    def _parse_ai_response(self, response_text: str) -> List[Dict[str, Any]]:
        """Parse Claude's JSON response into a list of question dicts."""

        response_text = response_text.strip()

        # Remove markdown code blocks if present
        if '```json' in response_text:
            match = re.search(r'```json\s*\n(.*?)\n```', response_text, re.DOTALL)
            if match:
                response_text = match.group(1).strip()
        elif '```' in response_text:
            response_text = re.sub(r'^```\s*\n', '', response_text)
            response_text = re.sub(r'\n```\s*$', '', response_text)
            response_text = response_text.strip()

        # Find JSON array if response has extra text around it
        if not response_text.startswith('['):
            json_match = re.search(r'\[.*\]', response_text, re.DOTALL)
            if json_match:
                response_text = json_match.group(0)

        try:
            questions = json.loads(response_text)

            if not isinstance(questions, list):
                return []

            validated = []
            for i, q in enumerate(questions):
                try:
                    if self._validate_question(q):
                        validated.append(q)
                    else:
                        logger.debug("Skipping question %d — failed validation", i + 1)
                except Exception as e:
                    logger.warning("Error validating question %d: %s", i + 1, e)
                    continue

            logger.debug("Found %d valid questions in AI response", len(validated))
            return validated

        except json.JSONDecodeError as e:
            logger.error("Failed to parse AI response: %s\nPreview: %s", e, response_text[:500])

            # Try to recover individual objects from a malformed response.
            # The pattern handles one level of nesting (needed for matching pairs
            # which contain inner {…} objects inside the matching_pairs array).
            try:
                objects = re.findall(
                    r'\{(?:[^{}]|\{[^{}]*\})*\}',
                    response_text,
                    re.DOTALL,
                )
                recovered = []
                for obj_str in objects:
                    try:
                        q = json.loads(obj_str)
                        if self._validate_question(q):
                            recovered.append(q)
                    except (ValueError, json.JSONDecodeError):
                        continue
                if recovered:
                    logger.warning("Recovered %d questions from malformed response", len(recovered))
                    return recovered
            except Exception as recover_err:
                logger.error("Recovery attempt failed: %s", recover_err)

            raise ValueError(f"Could not parse AI response: {str(e)}") from e

    def _validate_question(self, question: Dict[str, Any]) -> bool:
        """Validate a question has the minimum required fields"""

        # Must have type and question text at minimum
        if not question.get('type') or not question.get('question'):
            return False

        q_type = question['type']

        # Section headers are always valid — no further checks needed
        if q_type == 'section_header':
            question.setdefault('answer', '')
            question.setdefault('explanation', '')
            question.setdefault('difficulty', 'easy')
            question.setdefault('points', 0)
            return True

        # Multiple choice must have all 4 options
        if q_type == 'multiple_choice':
            for option in ['option_a', 'option_b', 'option_c', 'option_d']:
                if not question.get(option):
                    return False

        # Matching: ensure default arrays — do NOT silently drop the question
        if q_type == 'matching':
            if 'column_a' not in question or question.get('column_a') is None:
                question['column_a'] = []
            if 'column_b' not in question or question.get('column_b') is None:
                question['column_b'] = []
            if 'matching_pairs' not in question or question.get('matching_pairs') is None:
                question['matching_pairs'] = []

        # Set defaults for optional fields
        if question.get('difficulty') not in ['easy', 'medium', 'hard']:
            question['difficulty'] = 'medium'

        if not question.get('points'):
            question['points'] = 1

        if 'explanation' not in question:
            question['explanation'] = ''

        if 'answer' not in question:
            question['answer'] = ''

        return True


def get_extractor():
    """Factory function to get an extractor instance"""
    return AIQuestionExtractor()